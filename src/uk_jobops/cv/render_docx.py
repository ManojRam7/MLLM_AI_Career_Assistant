"""Render a TailoredCV into an ATS-safe single-column .docx (CV + cover letter)."""
from __future__ import annotations

import re
from pathlib import Path

from ..models import TailoredCV

BLUE = (0x1F, 0x4E, 0x79)


def _safe(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")[:60] or "job"


def _build(cfg_base_cv: dict, t: TailoredCV):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    blue = RGBColor(*BLUE)
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(2); n.paragraph_format.line_spacing = 1.0
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(1.0); s.left_margin = s.right_margin = Cm(1.5)

    def center(text, size=9.5, bold=False, color=None, italic=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color:
            r.font.color.rgb = color

    def heading(text):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = blue
        pPr = p._p.get_or_add_pPr(); bd = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "1F4E79")):
            bot.set(qn(k), v)
        bd.append(bot); pPr.append(bd)

    c = cfg_base_cv.get("contact", {})
    center(cfg_base_cv.get("name", ""), 22, bold=True)
    if t.job_title_line:
        center(t.job_title_line, 12, bold=True, color=blue)
    loc = c.get("location", "")
    if cfg_base_cv.get("contact", {}).get("open_to_relocation"):
        loc += " (open to relocation)"
    center(f"{loc}  |  {c.get('phone','')}  |  {c.get('email','')}")
    center(f"{c.get('linkedin','')}  |  {c.get('github','')}  |  {c.get('portfolio','')}")

    heading("Profile")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.add_run(t.profile)

    heading("Core Skills")
    for sk in t.skills:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.add_run((sk.get("label", "") + ": ")).bold = True; p.add_run(sk.get("items", ""))

    heading("Professional Experience")
    for e in t.experience:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(e.get("title", "")).bold = True; p.add_run("\t" + e.get("dates", ""))
        for b in e.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet"); bp.paragraph_format.space_after = Pt(2); bp.add_run(b)

    if t.projects:
        heading("Selected Projects")
        for pr in t.projects:
            bp = doc.add_paragraph(style="List Bullet"); bp.paragraph_format.space_after = Pt(2)
            bp.add_run(pr.get("name", "")).bold = True
            bp.add_run(" - " + pr.get("text", ""))

    heading("Education & Certifications")
    edu = t.education or "; ".join(f"{e['degree']} ({e['meta']})" for e in cfg_base_cv.get("education", []))
    doc.add_paragraph().add_run(edu)
    if cfg_base_cv.get("certifications"):
        p = doc.add_paragraph(); p.add_run("Certifications: ").bold = True; p.add_run(cfg_base_cv["certifications"])
    if cfg_base_cv.get("right_to_work"):
        p = doc.add_paragraph(); r = p.add_run("Right to work: " + cfg_base_cv["right_to_work"]); r.italic = True; r.font.size = Pt(9.5)
    return doc


def _cover(cfg_base_cv: dict, t: TailoredCV):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    s = doc.sections[0]; s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(1.8); s.left_margin = s.right_margin = Cm(2.2)
    c = cfg_base_cv.get("contact", {})
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(cfg_base_cv.get("name", "")).bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{c.get('phone','')} | {c.get('email','')} | {c.get('linkedin','')}").font.size = Pt(9.5)
    doc.add_paragraph()
    for line in t.cover_letter:
        para = doc.add_paragraph(line); para.paragraph_format.space_after = Pt(8); para.paragraph_format.line_spacing = 1.08
    return doc


def render(base_cv: dict, t: TailoredCV, job: dict, out_root: str = "output") -> tuple[str, str]:
    folder = Path(out_root) / f"{_safe(job.get('company',''))}_{_safe(job.get('title',''))}"
    folder.mkdir(parents=True, exist_ok=True)
    cv_path = folder / "CV.docx"
    cover_path = folder / "Cover_Letter.docx"
    _build(base_cv, t).save(cv_path)
    if t.cover_letter:
        _cover(base_cv, t).save(cover_path)
    return str(cv_path), str(cover_path if t.cover_letter else "")
